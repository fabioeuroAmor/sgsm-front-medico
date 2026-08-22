# -*- coding: utf-8 -*-
"""
Gera o docx final consolidado de evidencias de QA para IaPage (/ia).
Roda uma unica vez, le screenshots do TEMP e escreve o docx final no repo,
sobrescrevendo o arquivo da rodada anterior.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMP_DIR = os.path.join(os.environ["TEMP"], "qa-ia-integracao-2")
OUT_PATH = r"C:\AmbienteDev\sgsm-front-medico\docs\prd\ia-integracao\evidencias-2026-08-20.docx"

def img(name):
    p = os.path.join(TEMP_DIR, name)
    return p if os.path.exists(p) else None

doc = Document()

# ---- Title ----
title = doc.add_heading("QA Funcional — IaPage (/ia) — Integração sgsm-ia", level=0)

meta = doc.add_paragraph()
meta.add_run("Data da execução: 2026-08-20/21").bold = True
doc.add_paragraph(
    "Execução ao vivo via Playwright MCP contra http://localhost:3001, com backend real "
    "(sgsm-ia:8082, ms-sboot-auth:8081, sgsm core:8080), sem mocks nos itens funcionais. "
    "IA-24/25/26 usaram interceptação de rede real do Playwright (page.route) para simular "
    "indisponibilidade/latência do backend, sem derrubar nenhum serviço real."
)
doc.add_paragraph(
    "Credenciais de teste: paulo@gmail.com / qwertyui (perfil MEDICO). "
    "IA-01 foi reexecutado após correção do proxy do Vite (vite.config.ts) que capturava "
    "indevidamente a navegação da própria página /ia e /crm."
)

# ---- Summary ----
doc.add_heading("Resumo", level=1)
summary_table = doc.add_table(rows=1, cols=2)
summary_table.style = "Light Grid Accent 1"
hdr = summary_table.rows[0].cells
hdr[0].text = "Métrica"
hdr[1].text = "Valor"
rows_data = [
    ("Total de itens", "28 (IA-01 a IA-28)"),
    ("Aprovados", "28"),
    ("Reprovados", "0"),
    ("Não testável", "0"),
    ("Bugs críticos cross-cutting encontrados", "2 — ambos corrigidos e reconfirmados em 2026-08-21 (ver seção dedicada)"),
]
for k, v in rows_data:
    row = summary_table.add_row().cells
    row[0].text = k
    row[1].text = v

# ---- Cross-cutting findings ----
doc.add_heading("Achados críticos cross-cutting (não amarrados a um único item) — AMBOS CORRIGIDOS", level=1)
doc.add_paragraph(
    "Os dois achados abaixo motivaram a reprovação original de IA-16 e IA-19 em 2026-08-20. "
    "Ambos foram corrigidos em 2026-08-21 e reconfirmados ao vivo (ver itens IA-16/IA-19 "
    "revisados na seção de itens detalhados)."
)

doc.add_heading("1. Refresh token rotacionado nunca é persistido (causa raiz de IA-16) — CORRIGIDO", level=2)
p = doc.add_paragraph()
p.add_run("Arquivo: ").bold = True
p.add_run("src/services/api.ts, linha 38")
doc.add_paragraph(
    "O interceptor de resposta do axios, ao tratar um 401, chama POST /v1/api/auth/refresh "
    "e faz tokenStore.set(res.data.accessToken) — mas NUNCA persiste o campo refreshToken "
    "que o backend retorna no mesmo payload. Resposta real capturada durante o teste:"
)
code = doc.add_paragraph()
code.add_run(
    '{"accessToken":"eyJ...","expiresIn":900,'
    '"refreshToken":"8e6a8619-8c59-4710-8833-3cf77fa618ad","tipo":"Bearer"}'
).font.name = "Consolas"
doc.add_paragraph(
    "localStorage.refresh_token permaneceu com o valor ORIGINAL do login mesmo após duas "
    "chamadas de refresh bem-sucedidas (confirmado via localStorage.getItem('refresh_token') "
    "antes/depois). Como o backend rotaciona o refresh token a cada uso, o app volta a mandar "
    "um token já substituído na próxima chamada — reproduzido de forma consistente (5 de 5 "
    "tentativas nesta rodada): o SEGUNDO reload/hard-navigation em /ia após o login (ou após "
    "qualquer refresh bem-sucedido anterior) recebe 401 em /auth/refresh e força logout "
    "completo (redirect para /login), mesmo com o access token da sessão ainda dentro da "
    "validade de 15 minutos. Isso é a causa raiz da reprovação de IA-16."
)
p = doc.add_paragraph()
p.add_run("Correção aplicada: ").bold = True
p.add_run(
    "commit 317fc608 \"fix(auth): persiste o refreshToken rotacionado apos renovar sessao\" "
    "(sgsm-front-medico), mergeado em develop via qa/medicos. Reconfirmado ao vivo em "
    "2026-08-21: refresh_token rotaciona e persiste corretamente a cada reload, 3 reloads "
    "consecutivos sem logout indevido."
)

doc.add_heading("2. sgsm-ia retorna 403 (não 401) para chamadas não autenticadas (causa raiz de IA-19) — CORRIGIDO", level=2)
doc.add_paragraph(
    "O interceptor de resposta em api.ts só trata err.response?.status === 401 para disparar "
    "o fluxo de refresh-then-redirect-to-login. Nos testes, tanto um token corrompido quanto "
    "a ausência total do header Authorization em POST /ia/chat resultaram em 403 Forbidden "
    "do backend real (não 401 Unauthorized), diferente do padrão usado por ms-sboot-auth "
    "(login/refresh corretamente retornam 401 para credenciais inválidas). Como consequência, "
    "o código do interceptor destinado a tratar sessão expirada nunca é executado para "
    "/ia/chat — o usuário fica preso na página do assistente vendo apenas um erro genérico "
    "('Não consegui processar sua pergunta. Request failed with status code 403'), sem ser "
    "redirecionado ao /login e sem indicação clara de que precisa autenticar novamente."
)
p = doc.add_paragraph()
p.add_run("Correção aplicada: ").bold = True
p.add_run(
    "JwtAuthFilter.java (sgsm-ia) passou a chamar response.sendError(401) para token "
    "inválido/expirado/revogado, replicando o padrão já usado no sgsm core (fix 8dedb56, "
    "já em develop) — sem alterar o caso de token ausente, que segue o comportamento padrão "
    "do Spring Security (decisão deliberada e consistente entre os dois serviços). SecurityConfig "
    "do sgsm-ia liberou /error pelo mesmo motivo do sgsm core (o sendError dispara forward "
    "interno que reentraria na cadeia sem token). Reconfirmado ao vivo em 2026-08-21: token "
    "corrompido retorna 401, e o app redireciona corretamente para /login."
)

doc.add_heading("3. Nota metodológica (não é bug da aplicação)", level=2)
doc.add_paragraph(
    "Durante a simulação de IA-19, uma tentativa inicial usando monkey-patch de "
    "XMLHttpRequest.prototype.setRequestHeader deixou um patch residual ativo que corrompeu "
    "o header Authorization de requisições subsequentes reais (não simuladas), causando 403 "
    "também no primeiro envio de IA-22. Identificado e corrigido via reload+login limpo; "
    "IA-22 foi reexecutado com sucesso (200 OK) logo em seguida. Registrado aqui por "
    "transparência — não é um defeito do código da aplicação, é um artefato da própria "
    "metodologia de teste.")

# ---- Detailed items ----
doc.add_heading("Itens do plano de teste (IA-01 a IA-28)", level=1)

def add_item(item_id, title_text, status, desc_lines, evidence_lines, images):
    h = doc.add_heading(f"{item_id} — {title_text}", level=2)
    p = doc.add_paragraph()
    run = p.add_run(f"Status: {status}")
    run.bold = True
    if status.startswith("APROVADO"):
        run.font.color.rgb = RGBColor(0x1E, 0x7B, 0x34)
    elif status.startswith("REPROVADO"):
        run.font.color.rgb = RGBColor(0xB3, 0x1B, 0x1B)
    else:
        run.font.color.rgb = RGBColor(0x8A, 0x6D, 0x00)

    for line in desc_lines:
        doc.add_paragraph(line)

    if evidence_lines:
        p = doc.add_paragraph()
        p.add_run("Evidência (console/rede):").bold = True
        for line in evidence_lines:
            code_p = doc.add_paragraph()
            code_p.add_run(line).font.name = "Consolas"
            code_p.paragraph_format.space_after = Pt(2)

    for im in images:
        path = img(im)
        if path:
            doc.add_picture(path, width=Inches(6.0))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph(f"[screenshot não encontrado: {im}]")
    doc.add_paragraph("")


add_item(
    "IA-01",
    "Acessar /ia sem login → redireciona para /login",
    "APROVADO",
    [
        "Navegação direta para http://localhost:3001/ia sem sessão ativa (localStorage vazio) "
        "resultou em redirect imediato e limpo para /login, confirmando que a correção do "
        "proxy do Vite (regex restrito a /ia/(chat|busca|paciente|kpis|etl) e "
        "/crm/(leads|paciente|tags|churn)) resolveu o bug da rodada anterior.",
        "Nenhuma chamada de API foi disparada durante o redirect — é resolvido puramente client-side "
        "pela PrivateRoute checando ausência de token.",
    ],
    ["window.location.href = 'http://localhost:3001/login' (confirmado via page.url() e via "
     "localStorage vazio: '{}')"],
    ["IA-01_login_redirect.png"],
)

add_item(
    "IA-02",
    "Login válido + navegação via sidebar → Chat ativo com boas-vindas",
    "APROVADO",
    [
        "Login com paulo@gmail.com/qwertyui bem-sucedido, redirecionado para /pacientes. "
        "Clique em 'Assistente IA' na sidebar carregou /ia com a aba Chat ativa por padrão "
        "e a mensagem de boas-vindas do assistente já presente.",
    ],
    [
        "POST /v1/api/auth/login => 200 OK",
        "GET /v1/api/auth/me => 200 OK",
    ],
    ["IA-02_ia_page_chat_default.png"],
)

add_item(
    "IA-03",
    "Reload (F5) em /ia estando logado → continua autenticado",
    "APROVADO",
    [
        "Recarregar a página em /ia (primeiro reload após o login) manteve o usuário "
        "autenticado, sem voltar para /login.",
        "Nota: descoberto posteriormente (ver Achado crítico #1) que reloads SUBSEQUENTES "
        "falham de forma consistente devido a um bug real de rotação de refresh token. "
        "Este item testa especificamente o primeiro reload, que funciona.",
    ],
    [
        "POST /v1/api/auth/refresh => 200 OK (x2)",
        "GET /v1/api/auth/me => 200 OK (x2)",
    ],
    ["IA-03_reload_stays_authenticated.png"],
)

add_item(
    "IA-04",
    "Pergunta simples via botão enviar → resposta real do assistente",
    "APROVADO",
    [
        "Pergunta 'Quantos pacientes estão cadastrados?' enviada com o botão. Mensagem do "
        "usuário à direita, resposta real do backend RAG à esquerda: '5 pacientes estão "
        "cadastrados.' (valor consistente com a lista real de pacientes cadastrados).",
    ],
    [
        "POST /ia/chat => 200 OK, duration 4036ms",
        "Authorization: Bearer <jwt> presente no request header",
    ],
    ["IA-04_chat_response.png"],
)

add_item(
    "IA-05",
    "Enter (sem Shift) envia mensagem",
    "APROVADO",
    [
        "Pergunta 'Qual o total de médicos cadastrados?' enviada via tecla Enter. Mesmo "
        "comportamento do IA-04: resposta real 'O total de médicos cadastrados é 2.', sem "
        "quebra de linha indevida.",
    ],
    ["POST /ia/chat => 200 OK (x2, incluindo a mensagem anterior do IA-04)"],
    ["IA-05_enter_send.png"],
)

add_item(
    "IA-06",
    "Shift+Enter insere quebra de linha, não envia",
    "APROVADO",
    [
        "Digitado 'linha um', Shift+Enter, 'linha dois' via pressSequentially (digitação real "
        "caractere a caractere). Verificado via JS: textarea.value === 'linha um\\nlinha dois' "
        "— quebra de linha real inserida, nenhuma requisição disparada, nenhuma mensagem "
        "enviada ao histórico.",
    ],
    ["textarea.value = \"linha um\\nlinha dois\" (hasNewline: true)"],
    ["IA-06_shift_enter_newline.png"],
)

add_item(
    "IA-07",
    "Botão de enviar desabilitado com campo vazio ou só espaços",
    "APROVADO",
    [
        "Confirmado via interação real de teclado (Ctrl+A + Backspace para limpar, digitação "
        "de 3 espaços): botão enviar fica disabled=true tanto para campo vazio quanto para "
        "espaços em branco.",
        "Nota metodológica: uma tentativa inicial usando JS bruto (el.value = '' via DOM, "
        "sem passar pelo setter controlado do React) produziu uma leitura enganosa "
        "(disabled: false) porque não sincronizou corretamente o estado interno do React — "
        "não é um bug da aplicação, é uma limitação da técnica de teste. O reteste limpo via "
        "teclado real confirma o comportamento correto.",
    ],
    [
        "Campo vazio: textareaValue='\"\"', sendDisabled=true",
        "Campo com 3 espaços: textareaValue='\"   \"', sendDisabled=true",
    ],
    ["IA-07_empty_field_disabled.png"],
)

add_item(
    "IA-08",
    "Botão/campo bloqueados durante loading (evita envio duplicado)",
    "APROVADO",
    [
        "Imediatamente após o clique em enviar, o botão foi lido como disabled=true "
        "(programaticamente, via JS) — o backend responde rápido o suficiente (poucos "
        "segundos) que capturar visualmente o instante exato do spinner via screenshot foi "
        "inconsistente, mas o estado disabled foi confirmado no instante seguinte ao clique "
        "em múltiplas tentativas.",
    ],
    ["sendDisabled: true (checado programaticamente logo após o clique em enviar)"],
    ["IA-08_loading_state.png", "IA-08_loading_spinner2.png"],
)

add_item(
    "IA-09",
    "Auto-scroll até o final após receber resposta",
    "APROVADO",
    [
        "Viewport reduzido para forçar overflow. Após enviar mensagem e receber resposta, "
        "o container de mensagens (overflow-y-auto) teve seu scrollTop ajustado "
        "automaticamente até o final.",
    ],
    ["scrollTop=469, scrollHeight=724, clientHeight=255 → atBottom: true "
     "(|scrollHeight - clientHeight - scrollTop| < 5)"],
    ["IA-09_autoscroll.png"],
)

add_item(
    "IA-10",
    "Histórico mantido entre múltiplas perguntas",
    "APROVADO",
    [
        "Após várias perguntas e respostas em sequência (IA-04, IA-05, IA-08), todas as "
        "mensagens anteriores permaneceram visíveis na tela junto com as novas.",
    ],
    [],
    ["IA-10_history_preserved.png"],
)

add_item(
    "IA-11",
    "Pergunta de CRM/KPI com resposta multi-linha renderiza corretamente",
    "APROVADO",
    [
        "Pergunta 'Me dê um resumo dos indicadores de CRM disponíveis, um item por linha.' "
        "retornou uma resposta real longa e genuinamente multi-linha (9 itens), renderizada "
        "com quebras de linha reais na bolha (whitespace-pre-wrap funcionando), sem quebrar "
        "o layout.",
    ],
    ["POST /ia/chat => 200 OK; resposta com \\n reais entre itens, confirmados visualmente "
     "no screenshot (linhas separadas, não concatenadas)."],
    ["IA-11_multiline_response.png"],
)

add_item(
    "IA-12",
    "Aba KPIs carrega automaticamente e exibe indicadores",
    "APROVADO",
    [
        "Clique na aba KPIs disparou carregamento automático e exibiu os indicadores "
        "retornados por /ia/kpis em cards (pacientes alto valor, churn risco, ocupação de "
        "agenda, funil por médico, cancelamentos, faturamento mensal, resumo analítico).",
    ],
    ["GET /ia/kpis => 200 OK"],
    ["IA-12_kpis_load.png"],
)

add_item(
    "IA-13",
    "Alternar Chat→KPIs→Chat→KPIs não recarrega os dados",
    "APROVADO",
    [
        "Após o carregamento inicial dos KPIs, alternar para a aba Chat e voltar para KPIs "
        "não disparou uma nova requisição — os cards continuaram exibidos com os dados já "
        "carregados em estado (kpis), sem novo loading.",
    ],
    ["GET /ia/kpis: apenas 1 requisição (#85) mantida no log de rede após a troca de abas "
     "ida e volta."],
    ["IA-13_no_reload_on_tab_switch.png"],
)

add_item(
    "IA-14",
    "Botão 'Atualizar' dispara novo carregamento dos KPIs",
    "APROVADO",
    [
        "Clique em 'Atualizar' na aba KPIs disparou uma nova requisição e atualizou os cards.",
    ],
    ["GET /ia/kpis => 200 OK (nova requisição #86, distinta da #85 anterior)"],
    ["IA-14_atualizar_reload.png"],
)

add_item(
    "IA-15",
    "Valores objeto/array exibidos via JSON.stringify legível",
    "APROVADO",
    [
        "Confirmado na mesma evidência do IA-12: campos que são array/objeto (ocupacaoAgenda, "
        "funilMedico, resumo) aparecem como JSON formatado e legível "
        "(ex.: '[ { \"medico_id\": ..., \"medico_nome\": \"paulo souza\", ... } ]'), nunca "
        "como '[object Object]'. Campos vazios aparecem como '[]'.",
    ],
    [],
    ["IA-12_kpis_load.png"],
)

add_item(
    "IA-16",
    "Reload no meio de uma conversa → histórico perdido sem erro",
    "APROVADO (reteste 2026-08-21, após correção do Achado crítico #1)",
    [
        "Execução original (2026-08-20): REPROVADO — 5 de 5 reloads causavam logout indevido "
        "por causa do bug de rotação de refresh token (ver Achado crítico #1, já corrigido).",
        "Reteste (2026-08-21): login → /ia → 3 perguntas reais enviadas e respondidas pelo RAG "
        "→ 3 reloads reais consecutivos. Em NENHUM dos 3 o usuário foi deslogado — permaneceu "
        "autenticado em /ia (sidebar com 'paulo souza / MEDICO'), só o histórico do chat foi "
        "resetado para a mensagem de boas-vindas, que é o comportamento esperado. Console sem "
        "erros/warnings nos 3 reloads.",
    ],
    [
        "Reteste: POST /v1/api/auth/refresh => 200 OK (x3, um por reload)",
        "Reteste: GET /v1/api/auth/me => 200 OK (x3)",
        "Reteste: refresh_token em localStorage rotacionou e persistiu a cada reload "
        "(confirmado por prefixo distinto antes/depois de cada um)",
    ],
    ["IA-16_before.png", "IA-16_after.png", "IA16_03_after_reload1_still_authenticated.png",
     "IA16_04_after_reload2_still_authenticated.png", "IA16_05_after_reload3_still_authenticated.png"],
)

add_item(
    "IA-17",
    "Navegar para outra página e voltar → mesmo reset, sem erro",
    "APROVADO",
    [
        "Navegação via SPA (clique nos links da sidebar, sem reload de página) para "
        "/pacientes e de volta para /ia: sessão permaneceu autenticada (rota de navegação "
        "client-side não passa pelo fluxo de refresh, que só é acionado em hard reload), "
        "chat resetado para a mensagem de boas-vindas, sem erros.",
    ],
    ["Nenhuma chamada a /auth/refresh disparada (navegação puramente client-side)."],
    ["IA-17_before_navigate.png", "IA-17_after_navigate_back.png"],
)

add_item(
    "IA-18",
    "Clique duplo / Enter repetido não duplica envio",
    "APROVADO",
    [
        "Duplo-clique no botão de enviar com pergunta digitada resultou em apenas 1 "
        "requisição e 1 par de mensagens na tela (a segunda requisição observada no log de "
        "rede era residual de uma mensagem anterior, não uma duplicata).",
        "3 pressões rápidas de Enter em seguida também resultaram em apenas 1 nova "
        "requisição e 1 novo par de mensagens.",
    ],
    [
        "Duplo-clique: POST /ia/chat => 200 OK (apenas 1 nova requisição para a pergunta atual)",
        "Enter x3: POST /ia/chat => 200 OK (apenas 1 nova requisição para a pergunta atual)",
    ],
    ["IA-18_doubleclick_no_duplicate.png", "IA-18_rapid_enter_no_duplicate.png"],
)

add_item(
    "IA-19",
    "Sessão expirada ao enviar pergunta → redireciona para /login",
    "APROVADO (reteste 2026-08-21, após correção do Achado crítico #2)",
    [
        "Execução original (2026-08-20): REPROVADO — sgsm-ia retornava 403 (não 401) para "
        "token inválido, então o interceptor do front (que só trata 401) nunca redirecionava "
        "para /login (ver Achado crítico #2, já corrigido).",
        "Reteste (2026-08-21): refresh_token corrompido via localStorage + header Authorization "
        "da chamada de chat corrompido via patch single-shot em XMLHttpRequest.setRequestHeader "
        "(mecanismo real inspecionado em tokenStore.ts/api.ts antes do teste). Pergunta real "
        "enviada no chat.",
        "Resultado: POST /ia/chat retornou 401 (não mais 403) — a correção do JwtAuthFilter do "
        "sgsm-ia está funcionando. O refresh automático subsequente também falhou (401, "
        "refresh_token corrompido de propósito), e o app redirecionou corretamente para "
        "/login — tela de login renderizada normalmente, sem travamento.",
    ],
    [
        "Reteste: POST /ia/chat => 401 Unauthorized (era 403 na rodada anterior)",
        "Reteste: POST /v1/api/auth/refresh => 401 Unauthorized (refresh_token corrompido de propósito)",
        "Reteste: localStorage.refresh_token => null após o fluxo (limpo pelo interceptor)",
        "Reteste: location.href => http://localhost:3001/login",
        "Console: apenas 2x [ERROR] Failed to load resource 401 (logs padrão do navegador para "
        "resposta não-2xx, não exceção JS não tratada)",
    ],
    ["IA-19_session_expired_no_redirect.png", "IA19_01_before_corruption.png",
     "IA19_02_after_redirect_to_login.png"],
)

add_item(
    "IA-20",
    "ChatbotWidget global não bloqueia a interface do Assistente IA",
    "APROVADO",
    [
        "Ao abrir o ChatbotWidget flutuante (botão 'Assistente virtual', canto inferior "
        "direito), o painel do widget sobrepõe visualmente parte do histórico de mensagens "
        "do Assistente IA (comportamento esperado de qualquer overlay flutuante) — porém NÃO "
        "bloqueia nenhum elemento principal interativo: abas Chat/KPIs, campo de texto, "
        "botão de enviar e o próprio botão de fechar o widget permanecem visíveis e "
        "clicáveis. Fechado normalmente após o teste.",
    ],
    [],
    ["IA-20_widget_open.png"],
)

add_item(
    "IA-21",
    "Navegação por teclado (Tab) em ordem lógica com foco visível",
    "APROVADO",
    [
        "A partir do topo da página, Tab percorreu em ordem lógica: links da sidebar "
        "(Home → Pacientes → Médicos → Estabelecimentos → Serviços → Agendamentos → "
        "Funcionários → Assistente IA → CRM) → botão Sair → aba Chat → aba KPIs → região de "
        "scroll das mensagens (foco padrão do Chromium para containers com overflow-y:auto, "
        "comportamento normal do navegador) → textarea → botão enviar (habilitado após "
        "digitar texto). Indicador de foco visível (outline) em todos os elementos.",
    ],
    [],
    [
        "IA-21_focus_sidebar_link.png",
        "IA-21_focus_chat_tab.png",
        "IA-21_focus_textarea2.png",
        "IA-21_focus_send_button.png",
    ],
)

add_item(
    "IA-22",
    "Mensagem muito longa: textarea expande até maxHeight e rola internamente",
    "APROVADO",
    [
        "Mensagem de ~500 caracteres digitada: textarea expandiu com scroll interno "
        "(maxHeight: 120px confirmado via CSS computado, overflowY: auto, scrollHeight > "
        "clientHeight), sem quebrar o layout da página. Mensagem enviada e exibida "
        "corretamente na bolha, com quebra de linha (wrap) normal.",
        "Nota: a primeira tentativa de envio retornou 403 por causa do patch residual do "
        "XMLHttpRequest deixado ativo durante os testes de IA-19 (ver nota metodológica na "
        "seção de achados). Após relogin limpo, reenviada com sucesso (200 OK).",
    ],
    ["POST /ia/chat => 200 OK (reteste limpo)"],
    ["IA-22_long_message_textarea.png", "IA-22_long_message_final.png"],
)

add_item(
    "IA-23",
    "Caracteres especiais / HTML-like / XSS exibidos como texto puro",
    "APROVADO",
    [
        "Mensagem contendo <script>alert(1)</script>, emojis (😀🩺💉), acentuação "
        "(ção, ã, é, ü), <b>bold</b> e aspas simples/duplas foi enviada. O conteúdo foi "
        "renderizado literalmente como texto (React escapa por padrão) — nenhum script "
        "executado, nenhuma tag HTML interpretada, sem quebra de layout, emojis e acentos "
        "exibidos corretamente.",
    ],
    [],
    ["IA-23_special_chars_xss.png"],
)

add_item(
    "IA-24",
    "Backend do chat indisponível → mensagem de erro amigável + toast",
    "APROVADO",
    [
        "Interceptação real de rede via Playwright (page.route, sem derrubar nenhum serviço "
        "real) forçando POST /ia/chat a retornar 503. Resultado: bolha de erro amigável "
        "'Não consegui processar sua pergunta. Serviço indisponível (simulado para teste QA)' "
        "e toast 'Erro ao consultar o assistente' exibido no canto superior direito. "
        "Interface não travou, campo de texto voltou a ficar disponível.",
    ],
    ["POST /ia/chat => 503 Service Unavailable (simulado via page.route)"],
    ["IA-24_backend_error_chat.png", "IA-24_toast_immediate.png"],
)

add_item(
    "IA-25",
    "Rede lenta ao enviar pergunta → loading persiste até resposta/erro",
    "APROVADO",
    [
        "Latência real de 6 segundos injetada via page.route (delay + route.continue(), sem "
        "mock de conteúdo — a resposta real do backend foi entregue após o atraso). Duração "
        "real confirmada via inspeção da requisição de rede: 6015ms. O botão de enviar "
        "permaneceu desabilitado durante toda a espera, nenhuma requisição duplicada foi "
        "disparada, e a UI resolveu corretamente ao final sem travar.",
    ],
    ["POST /ia/chat => 200 OK, duration: 6015ms (delay real confirmado)",
     "Apenas 1 requisição por envio, mesmo com a espera prolongada."],
    ["IA-25_loading_spinner_caught.png"],
)

add_item(
    "IA-26",
    "Backend dos KPIs indisponível → erro tratado sem loading infinito",
    "APROVADO",
    [
        "Interceptação real de rede forçando GET /ia/kpis a retornar 503. Resultado: toast "
        "'Erro ao carregar KPIs' exibido, painel de indicadores mostrou 'Nenhum dado "
        "carregado' (estado tratado explicitamente), sem cards quebrados nem loading "
        "infinito.",
    ],
    ["GET /ia/kpis => 503 Service Unavailable (simulado via page.route)"],
    ["IA-26_kpis_error.png"],
)

add_item(
    "IA-27",
    "Console limpo durante todo o fluxo (login → chat → KPIs → erro simulado)",
    "APROVADO",
    [
        "Revisão de todas as mensagens de console (nível error/warning) acumuladas durante "
        "toda a sessão de testes. Todas as entradas de erro encontradas são falhas de rede "
        "esperadas (401/403/503) originadas de: (a) simulações deliberadas dos itens IA-19, "
        "IA-24, IA-25, IA-26; (b) reproduções do bug de sessão do Achado crítico #1 durante "
        "a investigação do IA-16. Nenhuma exceção JS não tratada, erro de React "
        "(ex.: 'Warning: ...', 'Uncaught TypeError') ou promise rejeitada sem tratamento foi "
        "encontrada em nenhum momento do fluxo.",
        "Observação à parte (não relacionada ao código da aplicação): mensagens de "
        "WebSocket 'ws://localhost:3001/' ERR_CONNECTION_REFUSED do cliente HMR do Vite "
        "apareceram esporadicamente — ruído de infraestrutura do servidor de "
        "desenvolvimento, não do código React.",
    ],
    [],
    [],
)

add_item(
    "IA-28",
    "/ia/chat e /ia/kpis chamados com header Authorization: Bearer <token>",
    "APROVADO",
    [
        "Confirmado em múltiplas requisições reais capturadas ao longo de toda a execução: "
        "tanto POST /ia/chat quanto GET /ia/kpis sempre incluem o header "
        "'authorization: Bearer <jwt>' via interceptor de request do api.ts.",
    ],
    [
        "POST /ia/chat → authorization: Bearer eyJhbGciOiJIUzUxMiJ9... (confirmado em "
        "múltiplas chamadas)",
        "GET /ia/kpis → authorization: Bearer eyJhbGciOiJIUzUxMiJ9... (confirmado)",
    ],
    ["IA-28_auth_header_confirmed.png"],
)

doc.save(OUT_PATH)
print("Saved:", OUT_PATH)
