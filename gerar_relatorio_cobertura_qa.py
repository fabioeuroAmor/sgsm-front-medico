from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Margens ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_color(run, hex_color):
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    set_color(run, "0f172a")
    return p

def subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    set_color(run, "64748b")
    return p

def heading1(text, color="1a56db"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    set_color(run, color)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_color(run, "1e293b")
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_color(run, "1e293b")
    return p

def note(text, color="b45309", bg=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    set_color(run, color)
    return p

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, accent="1a56db"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        set_color(run, "ffffff")
        shade_cell(hdr_cells[i], accent)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            set_color(run, "1e293b")
            if r_i % 2 == 1:
                shade_cell(cells[i], "f1f5f9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

# ── Capa ──────────────────────────────────────────────────────────────────────
title("Cobertura de QA Funcional — Front-end")
subtitle("sgsm-front-medico  ·  Skill executa-testes-funcionais-qa-front-end  ·  19/08/2026")

body(
    "Este relatório consolida quais telas do sistema já passaram por execução de QA funcional "
    "ao vivo (via Playwright, contra a aplicação e os back-ends reais, sem mocks) e quais telas "
    "ainda não têm nenhuma cobertura registrada. A fonte dos dados são os Pull Requests fechados "
    "no repositório sgsm-front-medico e as branches qa/* existentes no remoto até a data acima."
)

# ── Casos de uso testados ───────────────────────────────────────────────────────
heading1("Casos de uso testados")

add_table(
    ["Página / Rota", "Branch(es)", "Casos executados", "Resultado", "PR"],
    [
        ["/pacientes", "qa/pacientes-validacao*\nqa/pacientes-regressao",
         "70 (validação de campos)\n52/52 (regressão geral)", "Aprovado", "#14 (mergiado)"],
        ["/medicos", "qa/medicos",
         "57 (52 aprovados, 4 bloqueados\npor bug de backend, 1 não-testável)",
         "Aprovado com pendências\nde backend", "Sem PR\n(merge direto p/ develop)"],
        ["/estabelecimentos", "qa/estabelecimento-validacao",
         "54 (48 aprovados na 1ª rodada;\n100% após correções e regressão)", "Aprovado", "#15 (mergiado)"],
        ["/crm", "qa/crm", "74 (68 aprovados na rodada final)", "Aprovado", "#17 (mergiado)"],
        ["/servicos", "qa/servico-medico",
         "65 (61 aprovados, 3 bloqueados\npor limitação de ambiente,\n1 falha conhecida fora de escopo)",
         "Aprovado", "#19 (mergiado)"],
    ],
    col_widths=[3.2, 3.3, 4.3, 3.0, 3.2],
)
note("* qa/pacientes-validacao não corresponde a uma branch remota própria; o test-plan de validação de campos foi executado e documentado dentro do histórico que resultou na branch qa/pacientes-regressao.")

# ── Casos de uso faltando ────────────────────────────────────────────────────────
heading1("Casos de uso faltando")

add_table(
    ["Página / Rota", "Componente", "Observação"],
    [
        ["/", "HomePage", "Dashboard inicial — sem QA funcional registrada"],
        ["/world", "WorldPage", "Landing/apresentação — sem QA funcional registrada"],
        ["/login", "LoginPage", "Porta de entrada do sistema — nunca testada"],
        ["/registrar", "RegisterPage", "Fluxo de autocadastro — nunca testado"],
        ["/esqueci-senha", "EsqueciSenhaPage", "Fluxo de recuperação de senha — nunca testado"],
        ["/resetar-senha", "ResetarSenhaPage", "Fluxo de recuperação de senha — nunca testado"],
        ["/agendamentos", "AgendamentosPage", "Agenda médica, horários e conflitos — sem cobertura"],
        ["/funcionarios", "FuncionariosPage", "Sem cobertura"],
        ["/ia", "IaPage", "Integração com sgsm-ia — sem cobertura"],
    ],
    col_widths=[3.4, 3.6, 10.0],
)

# ── Bugs de backend — qa/medicos ────────────────────────────────────────────────
heading1("Bugs de backend identificados no QA de Médicos")

body(
    "Os 4 itens reprovados na QA de Médicos (qa/medicos) e atribuídos a backend já foram todos "
    "corrigidos e mergiados na develop, tanto no sgsm quanto no sgsm-front-medico."
)

add_table(
    ["Item", "Causa raiz", "Correção", "Status"],
    [
        ["M15 / M16 / M18\n(filtros ignorados)",
         "MedicoService.listar() restringia perfil MEDICO a ver\nsó o próprio cadastro, ignorando ativo/especialidade",
         "Removida a auto-restrição\nem MedicoService (sgsm)",
         "Corrigido\nsgsm PR #13 (0ef000b)"],
        ["M55\n(médicos recém-criados somem da lista)",
         "Mesma causa raiz de M15/M16/M18",
         "Resolvido pela mesma\ncorreção acima",
         "Corrigido\nsgsm PR #13 (0ef000b)"],
        ["M54\n(logout no 2º F5)",
         "Não era bug de backend: o sgsm-auth já rotacionava\no refresh token corretamente. O front-end\n(api.ts/useAuth.tsx) não persistia o token rotacionado",
         "Corrigido no front-end,\ncombinado com a correção\nde concorrência do qa/medicos",
         "Corrigido\nsgsm-front-medico (317fc60, 1337ddf)"],
    ],
    col_widths=[3.6, 6.5, 4.5, 4.4],
)
note("Conclusão: nenhuma pendência de backend restante do qa/medicos.", color="15803d")

# ── Observações e recomendações ─────────────────────────────────────────────────
heading1("Observações e recomendações")

bullet("O fluxo completo de autenticação (login, registro, esqueci/resetar senha) nunca passou por QA funcional dedicada — é o maior risco identificado, por ser a porta de entrada de todo o sistema.")
bullet("Agendamentos é provavelmente o módulo mais crítico ainda sem cobertura, por envolver agenda médica, horários e conflitos de marcação.")
bullet("qa/servico-medico foi mergiada na develop (PR #19) — os 5 módulos com QA funcional executada (Pacientes, Médicos, Estabelecimentos, CRM, Serviços) estão todos na develop.")
bullet("A tela de Médicos (qa/medicos) foi mergiada diretamente na develop, sem passar por Pull Request no GitHub — fora do padrão observado nas demais QAs.")
bullet("Os bugs de backend do qa/medicos (M15, M16, M18, M55, M54) já foram todos corrigidos e mergiados — ver seção dedicada acima. O achado de Serviços (limpar Duração em edição não persiste) segue como limitação conhecida e deliberadamente não corrigida, por ser convenção global do backend (ModelMapper Conditions.isNotNull()) que afeta todos os endpoints de update do sistema, não só Serviços.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Gerado a partir do histórico de PRs e branches qa/* do repositório sgsm-front-medico.")
run.font.size = Pt(9)
run.font.italic = True
set_color(run, "94a3b8")

doc.save("docs/prd/cobertura-qa/relatorio-cobertura-qa-front-end-2026-08-19.docx")
print("OK")
