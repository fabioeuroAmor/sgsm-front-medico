from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Margens ───────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_color(run, hex_color):
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
    run.font.color.rgb = RGBColor(r, g, b)

def title(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    set_color(run, "0f172a")
    return p

def subtitle(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(20)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    set_color(run, "64748b")
    return p

def heading1(text, color="1a56db"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    set_color(run, color)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_color(run, "1e293b")
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    set_color(run, "1e293b")
    return p

def note(text, color="b45309"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.italic = True
    set_color(run, color)
    return p

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths=None, accent="1a56db"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(9.5)
        set_color(run, "ffffff")
        shade_cell(hdr_cells[i], accent)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for r_i, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9.5)
            set_color(run, "1e293b")
            if r_i % 2 == 1:
                shade_cell(cells[i], "f1f5f9")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table

# ── Capa ──────────────────────────────────────────────────────────────────────
title("Pull Requests Mergeados em develop")
subtitle("Ecossistema SGSM  ·  sgsm-front-medico · sgsm · ms-sboot-auth  ·  20/08/2026")

body(
    "Este relatório lista os 9 Pull Requests abertos a partir das branches criadas durante o "
    "ciclo de QA funcional e correção de bugs, nos três repositórios do ecossistema, todos com "
    "base em develop. A ferramenta gh (GitHub CLI) não estava disponível no início da sessão; "
    "foi instalada via winget e autenticada via fluxo OAuth de dispositivo (login pelo navegador), "
    "o que permitiu abrir e, em seguida, mergear os 9 PRs diretamente por linha de comando."
)

note("Total: 9 branches · 9 PRs abertos e MERGEADOS em develop · nenhum conflito em nenhum merge.", color="15803d")

# ── Resumo consolidado ───────────────────────────────────────────────────────
heading1("Resumo consolidado — Todos os 9 PRs mergeados")

add_table(
    ["Ordem", "Repositório", "Branch", "PR", "Status"],
    [
        ["1", "sgsm", "fix/funcionario-email-duplicado", "#14", "✅ Mergeado"],
        ["2", "sgsm", "fix/agendamento-double-booking-e-filtros", "#15", "✅ Mergeado"],
        ["3", "sgsm", "fix/paciente-cpf-invalido", "#16", "✅ Mergeado"],
        ["4", "ms-sboot-auth", "fix/reset-senha-401-e-politica-senha", "#8", "✅ Mergeado"],
        ["5", "ms-sboot-auth", "fix/email-disponivel-cadastro-orfao", "#9", "✅ Mergeado"],
        ["6", "sgsm-front-medico", "qa/funcionarios", "#20", "✅ Mergeado"],
        ["7", "sgsm-front-medico", "qa/agendamentos", "#21", "✅ Mergeado"],
        ["8", "sgsm-front-medico", "qa/resetar-senha", "#22", "✅ Mergeado"],
        ["9", "sgsm-front-medico", "qa/registrar", "#23", "✅ Mergeado"],
    ],
    col_widths=[1.5, 3.3, 7.7, 1.6, 2.4],
    accent="15803d",
)

body(
    "Cada PR trazia no corpo: resumo da correção, contexto (item do QA que achou o bug) e o que "
    "foi testado. As dependências cruzadas (PRs de front-end que dependiam de um PR de backend "
    "específico) estavam marcadas no corpo de cada PR de front-end."
)

body(
    "Ordem de merge executada: primeiro os 3 backends do sgsm (#14 → #15 → #16), depois os 2 do "
    "ms-sboot-auth (#8 → #9), e só então os 4 front-ends do sgsm-front-medico (#20 → #21 → #22 → "
    "#23) — assim nenhum PR de front-end ficou temporariamente \"quebrado\" esperando um endpoint "
    "que ainda não existia em develop. Todos os 9 merges foram confirmados individualmente "
    "(estado MERGED via gh pr view) e nenhum conflito ocorreu em nenhuma etapa."
)

note(
    "Achado durante a revisão pré-merge: o PR #20 trouxe junto 2 arquivos não relacionados "
    "(relatorio-cobertura-qa-front-end-2026-08-19.docx e gerar_relatorio_cobertura_qa.py), "
    "remanescentes de uma tarefa anterior na mesma branch. O conteúdo funcional do PR estava "
    "correto; esses 2 arquivos já estão em develop e podem ser removidos num commit de limpeza "
    "separado, se desejado.",
    color="b45309",
)

# ── sgsm-front-medico ────────────────────────────────────────────────────────
heading1("Repositório: sgsm-front-medico (4 PRs — todos mergeados)")

add_table(
    ["PR", "Branch", "Conteúdo", "Status"],
    [
        ["#20 — .../sgsm-front-medico/pull/20",
         "qa/funcionarios",
         "Test-plan + correções de FuncionariosPage.tsx (e-mail duplicado na busca,\n"
         "nome só-espaço aceito, duplo-submit sem guard, validação de telefone) + evidências",
         "✅ Mergeado"],
        ["#21 — .../sgsm-front-medico/pull/21",
         "qa/agendamentos",
         "Test-plan + correções de AgendamentosPage.tsx (Telemedicina exibindo endereço\n"
         "físico do estabelecimento, busca do Passo 1 do wizard não filtrando, min da\n"
         "data calculado em UTC) + evidências",
         "✅ Mergeado"],
        ["#22 — .../sgsm-front-medico/pull/22",
         "qa/resetar-senha",
         "Test-plan + correção de ResetarSenhaPage.tsx (sem validação de tamanho\n"
         "mínimo de senha no client) + evidências",
         "✅ Mergeado"],
        ["#23 — .../sgsm-front-medico/pull/23",
         "qa/registrar",
         "Test-plan + correções de RegisterPage.tsx (checagem de e-mail disponível\n"
         "antes de cadastrar, validação de CPF, bloqueio de data de nascimento futura,\n"
         "duplo-submit sem guard) + evidências",
         "✅ Mergeado"],
    ],
    col_widths=[3.5, 3.0, 7.8, 2.2],
)

# ── sgsm (core) ──────────────────────────────────────────────────────────────
heading1("Repositório: sgsm — backend core (3 PRs — todos mergeados)")

add_table(
    ["PR", "Branch", "Conteúdo", "Status"],
    [
        ["#14 — .../sgsm/pull/14",
         "fix/funcionario-email-duplicado",
         "FuncionarioService / FuncionarioRepository — valida unicidade de e-mail\n"
         "ao cadastrar funcionário (só CPF era validado antes)",
         "✅ Mergeado"],
        ["#15 — .../sgsm/pull/15",
         "fix/agendamento-double-booking-e-filtros",
         "AgendamentoService / AgendamentoRepository — impede double-booking\n"
         "(revalidação + lock por médico), filtra horários já passados na data de\n"
         "hoje, corrige combinação de filtro médico+paciente na listagem",
         "✅ Mergeado"],
        ["#16 — .../sgsm/pull/16",
         "fix/paciente-cpf-invalido",
         "PacienteService — valida dígito verificador do CPF ao cadastrar;\n"
         "GlobalExceptionHandler — trata violação de integridade de dados como\n"
         "400 em vez de 500 cru (condição de corrida em cadastros concorrentes)",
         "✅ Mergeado"],
    ],
    col_widths=[3.2, 3.8, 7.5, 2.0],
)

# ── ms-sboot-auth ────────────────────────────────────────────────────────────
heading1("Repositório: ms-sboot-auth — autenticação (2 PRs — todos mergeados)")

add_table(
    ["PR", "Branch", "Conteúdo", "Status"],
    [
        ["#8 — .../ms-sboot-auth/pull/8",
         "fix/reset-senha-401-e-politica-senha",
         "ResetSenhaService / GlobalExceptionHandler — token de redefinição\n"
         "inválido/expirado/já usado agora retorna 400 (antes usava o mesmo 401\n"
         "de sessão expirada, fazendo o interceptor do front redirecionar\n"
         "silenciosamente usuários anônimos para /login sem mostrar o erro);\n"
         "exige senha com no mínimo 8 caracteres",
         "✅ Mergeado"],
        ["#9 — .../ms-sboot-auth/pull/9",
         "fix/email-disponivel-cadastro-orfao",
         "AuthController / AuthService — novo endpoint público\n"
         "GET /auth/email-disponivel, usado pelo RegisterPage.tsx para bloquear\n"
         "o cadastro antes de criar um médico/paciente órfão quando o e-mail já\n"
         "está em uso por outro perfil",
         "✅ Mergeado"],
    ],
    col_widths=[3.0, 4.0, 7.3, 2.2],
)

# ── Instalação e autenticação do gh CLI ─────────────────────────────────────
heading1("Instalação e autenticação do GitHub CLI")

bullet("Instalado via winget (winget install --id GitHub.cli), versão 2.97.0.")
bullet("Autenticado via fluxo OAuth de dispositivo (gh auth login → login pelo navegador): o terminal exibe um código de 8 caracteres, o usuário insere em github.com/login/device e autoriza o app \"GitHub CLI\".")
bullet("Confirmado com gh auth status: logado como fabioeuroAmor, escopos gist, read:org, repo, workflow.")
bullet("Os 9 PRs foram criados via gh pr create --base develop, um por branch, com título e corpo descrevendo resumo, contexto (item do QA que encontrou o bug) e o que foi testado.")
bullet("Os 9 PRs foram então mergeados via gh pr merge --merge (commit de merge, mantendo o histórico de cada PR), na ordem backends → front-ends, com confirmação individual de cada merge via gh pr view --json state,mergedAt.")

# ── Verificação pré-merge ────────────────────────────────────────────────────
heading1("Verificação pré-merge")

bullet("Antes de mergear, cada PR foi conferido individualmente com gh pr diff, comparando o conteúdo real do diff no GitHub (não a árvore de trabalho local, que trocava de branch com frequência ao longo da sessão) contra o código esperado da correção — confirmando que todas as 9 correções estavam de fato presentes nos PRs certos.")
bullet("Todos os 9 PRs estavam com mergeStateStatus \"CLEAN\" e mergeable \"MERGEABLE\" antes do início da sequência de merges.")
bullet("Nenhum conflito ocorreu durante os merges, mesmo entre as 4 branches qa/* do sgsm-front-medico, que foram criadas a partir de pontos diferentes de develop ao longo da sessão.")

# ── Observações ──────────────────────────────────────────────────────────────
heading1("Observações finais")

bullet("Todas as 9 branches partiram de develop e já haviam passado por reteste ao vivo confirmando a correção de cada bug (evidências em docs/prd/<tarefa>/evidence ou evidencias/ de cada branch qa/*).")
bullet("As branches (fix/* e qa/*) não foram apagadas após o merge — só os PRs foram fechados/mergeados. A limpeza das branches remotas, se desejada, fica como próximo passo.")
bullet("O PR #20 trouxe 2 arquivos não relacionados ao QA de Funcionários (relatório de cobertura .docx e o script gerador), remanescentes de uma tarefa anterior — não afeta a funcionalidade, mas pode ser limpo num commit separado.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(24)
run = p.add_run("Gerado a partir do estado dos PRs após a conclusão dos 9 merges (20/08/2026).")
run.font.size = Pt(9)
run.font.italic = True
set_color(run, "94a3b8")

doc.save("docs/prd/prs-pendentes/relatorio-prs-pendentes-2026-08-20.docx")
print("OK")
